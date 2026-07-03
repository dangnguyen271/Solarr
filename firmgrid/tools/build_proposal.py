"""Render the proposal to DOCX (python-docx) and A4 HTML (for Chrome -> PDF).

Run:  python build_proposal.py
Outputs land in ../submission/.
"""

import html
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from proposal_content import MEMBERS, PROJECT, SECTIONS, TEAM_NAME, TITLE

OUT = Path(__file__).resolve().parent.parent / "submission"
OUT.mkdir(exist_ok=True)
BASE = f"{TEAM_NAME}_{PROJECT}_PreliminaryRound"

NAVY = RGBColor(0x0F, 0x2A, 0x4A)
GREEN = RGBColor(0x1E, 0x7A, 0x46)


# --------------------------------------------------------------------- #
# DOCX
# --------------------------------------------------------------------- #
def build_docx() -> Path:
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Cm(1.6)
        sec.left_margin = sec.right_margin = Cm(1.8)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.18

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("ASIAN HACKATHON FOR GREEN FUTURE 2026 — PROJECT PROPOSAL")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = NAVY

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = GREEN

    p = doc.add_paragraph()
    r = p.add_run(f"TEAM NAME: {TEAM_NAME}")
    r.bold = True

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, head in enumerate(("No.", "Full Name", "Role")):
        cell = table.rows[0].cells[i]
        cell.text = head
        cell.paragraphs[0].runs[0].bold = True
    for no, name, role in MEMBERS:
        row = table.add_row()
        row.cells[0].text = no
        row.cells[1].text = name
        row.cells[2].text = role

    for sec_title, blocks in SECTIONS:
        hp = doc.add_paragraph()
        hp.paragraph_format.space_before = Pt(12)
        hp.paragraph_format.space_after = Pt(4)
        r = hp.add_run(sec_title)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = NAVY
        for kind, payload in blocks:
            if kind == "p":
                para = doc.add_paragraph(payload)
                para.paragraph_format.space_after = Pt(6)
            elif kind == "bullets":
                for b in payload:
                    para = doc.add_paragraph(b, style="List Bullet")
                    para.paragraph_format.space_after = Pt(3)

    path = OUT / f"{BASE}.docx"
    doc.save(path)
    return path


# --------------------------------------------------------------------- #
# HTML -> PDF (Chrome headless)
# --------------------------------------------------------------------- #
CSS = """
@page { size: A4; margin: 13mm 13mm; }
body { font-family: 'Helvetica Neue', Arial, sans-serif; font-size: 9.2pt;
       color: #1f2937; line-height: 1.44; }
.head { text-align: center; font-weight: 700; font-size: 12.5pt; color: #0f2a4a;
        letter-spacing: 0.4px; }
.title { text-align: center; font-weight: 700; font-size: 10.5pt; color: #1e7a46;
         margin: 4px 0 9px; line-height: 1.35; }
h2 { font-size: 10.3pt; color: #0f2a4a; margin: 10px 0 6px;
     border-left: 4px solid #1e7a46; background: #f2f7f3;
     padding: 2.5px 8px; border-radius: 2px; letter-spacing: 0.3px;
     page-break-after: avoid; }
p { margin: 0 0 6px; text-align: left; }
ul { margin: 0 0 7px 15px; padding: 0; }
li { margin-bottom: 4px; text-align: left; }
table { border-collapse: collapse; width: 100%; margin: 4px 0 8px; }
td, th { border: 1px solid #b6c2d2; padding: 3px 8px; font-size: 9.2pt; text-align: left; }
th { background: #eef4ee; }
.team { font-weight: 700; margin-top: 7px; }
.refs ul { column-count: 2; column-gap: 6mm; margin-left: 13px; }
.refs li { font-size: 8.1pt; line-height: 1.35; margin-bottom: 3.5px;
           break-inside: avoid; color: #374151; }
"""


def build_html_pdf() -> Path:
    parts = [
        "<meta charset='utf-8'>",
        f"<style>{CSS}</style>",
        "<div class='head'>ASIAN HACKATHON FOR GREEN FUTURE 2026 — PROJECT PROPOSAL</div>",
        f"<div class='title'>{html.escape(TITLE)}</div>",
        f"<p class='team'>TEAM NAME: {html.escape(TEAM_NAME)}</p>",
        "<table><tr><th>No.</th><th>Full Name</th><th>Role</th></tr>",
    ]
    for no, name, role in MEMBERS:
        parts.append(
            f"<tr><td>{no}</td><td>{html.escape(name)}</td><td>{html.escape(role)}</td></tr>"
        )
    parts.append("</table>")

    for sec_title, blocks in SECTIONS:
        is_refs = sec_title.startswith("8.")
        if is_refs:
            parts.append("<div class='refs'>")
        parts.append(f"<h2>{html.escape(sec_title)}</h2>")
        for kind, payload in blocks:
            if kind == "p":
                parts.append(f"<p>{html.escape(payload)}</p>")
            elif kind == "bullets":
                parts.append(
                    "<ul>" + "".join(f"<li>{html.escape(b)}</li>" for b in payload) + "</ul>"
                )
        if is_refs:
            parts.append("</div>")

    html_path = OUT / f"{BASE}.html"
    html_path.write_text("\n".join(parts), encoding="utf-8")

    pdf_path = OUT / f"{BASE}.pdf"
    chrome = "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome"
    subprocess.run(
        [
            chrome, "--headless", "--disable-gpu", "--no-pdf-header-footer",
            f"--print-to-pdf={pdf_path}", html_path.as_uri(),
        ],
        check=True,
        capture_output=True,
        timeout=120,
    )
    html_path.unlink()  # intermediate only
    return pdf_path


if __name__ == "__main__":
    words = 0
    for _, blocks in SECTIONS:
        for kind, payload in blocks:
            if kind == "p":
                words += len(payload.split())
            elif kind == "bullets":
                words += sum(len(b.split()) for b in payload)
    print(f"[content] ~{words} words")
    print(f"[docx] {build_docx()}")
    print(f"[pdf ] {build_html_pdf()}")
